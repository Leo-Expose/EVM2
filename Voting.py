import serial
import time
import os
from flask import Flask, jsonify, render_template
from threading import Thread
import tkinter as tk
from tkinter import messagebox
from docx import Document

# Define the serial port and baud rate.
SERIAL_PORT = 'COM4'  # Update this to the correct port
BAUD_RATE = 9600
BACKUP_FILE = os.path.expanduser("~/Desktop/Voting/votes_backup.txt")
final_results_file = os.path.expanduser("~/Desktop/Voting/final_results.docx")

# Candidate names
candidates = ["Candidate 1", "Candidate 2", "Candidate 3", "Candidate 4", "Candidate 5", "Candidate 6"]
votes = {candidate: 0 for candidate in candidates}
last_vote = None

# Load existing votes from backup file
if os.path.exists(BACKUP_FILE):
    with open(BACKUP_FILE, 'r') as f:
        for line in f:
            name, vote_count = line.strip().split(',')
            votes[name] = int(vote_count)

# Initialize serial connection
ser = serial.Serial(SERIAL_PORT, BAUD_RATE)

def save_backup():
    with open(BACKUP_FILE, 'w') as f:
        for name, count in votes.items():
            f.write(f"{name},{count}\n")

def export_results():
    doc = Document()
    doc.add_heading('Final Voting Results', 0)
    doc.add_heading('Post: Head Boy', level=1)
    
    # Find the winner
    winner = max(votes, key=votes.get)
    doc.add_paragraph(f"Winner: {winner} with {votes[winner]} votes")
    
    for name, count in votes.items():
        doc.add_paragraph(f"{name}: {count} votes")
    doc.save(final_results_file)
    print("Results saved to", final_results_file)

def print_results():
    result_file = os.path.expanduser("~/Desktop/Voting/results.txt")
    with open(result_file, 'w') as f:
        for name, count in votes.items():
            f.write(f"{name}: {count} votes\n")
    print("Results saved to", result_file)

def stop_voting():
    global running
    running = False
    messagebox.showinfo("Voting Stopped", "Voting has been stopped and results exported.")
    export_results()
    root.destroy()

# Flask app setup
app = Flask(__name__)

@app.route('/votes')
def get_votes():
    global last_vote
    return render_template('Live Tally.html', votes=votes, last_vote=last_vote)

@app.route('/latest_vote')
def get_latest_vote():
    global last_vote
    return render_template('Latest Vote.html', last_vote=last_vote)

def start_flask():
    app.run(debug=True, use_reloader=False, host='192.168.1.1', port=4200)

def start_latest_vote_server():
    latest_vote_app = Flask(__name__)

    @latest_vote_app.route('/latest_vote')
    def latest_vote():
        global last_vote
        return render_template('Latest Vote.html', last_vote=last_vote)
    
    latest_vote_app.run(debug=True, use_reloader=False, host='192.168.1.1', port=7000)

def start_voting():
    global running
    global last_vote
    running = True
    try:
        while running:
            if ser.in_waiting > 0:
                vote = ser.readline().decode().strip()
                if vote.startswith("Vote: "):
                    candidate_index = int(vote.split(" ")[1])
                    candidate_name = candidates[candidate_index]
                    votes[candidate_name] += 1
                    last_vote = candidate_name
                    print(f"{candidate_name} received a vote. Total now: {votes[candidate_name]}")
                    save_backup()
            time.sleep(0.1)
    except KeyboardInterrupt:
        print("Vote counting stopped.")
        print_results()
    finally:
        ser.close()

def set_running(value):
    global running
    running = value
    if running:
        voting_thread = Thread(target=start_voting)
        voting_thread.start()

# Start Flask servers in separate threads
flask_thread = Thread(target=start_flask)
flask_thread.start()

latest_vote_thread = Thread(target=start_latest_vote_server)
latest_vote_thread.start()

# Tkinter GUI setup
root = tk.Tk()
root.title("Voting Control")

start_button = tk.Button(root, text="Start Voting", command=lambda: set_running(True), font=('Helvetica', 16))
start_button.pack(pady=10)

pause_button = tk.Button(root, text="Pause Voting", command=lambda: set_running(False), font=('Helvetica', 16))
pause_button.pack(pady=10)

stop_button = tk.Button(root, text="Stop Voting", command=stop_voting, font=('Helvetica', 16))
stop_button.pack(pady=10)

root.geometry("300x200")
root.protocol("WM_DELETE_WINDOW", stop_voting)

# Start the Tkinter main loop
root.mainloop()
